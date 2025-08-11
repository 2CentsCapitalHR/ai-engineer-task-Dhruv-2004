from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import RGBColor


def _find_paragraph_index(document: Document, needle: str) -> Optional[int]:
    if not needle:
        return None
    target = needle.strip().lower()
    for idx, p in enumerate(document.paragraphs):
        if target in (p.text or "").lower():
            return idx
    return None


def _heuristic_keywords(issue_text: str) -> List[str]:
    text = (issue_text or "").lower()
    keywords: List[str] = []
    if any(k in text for k in ["jurisdiction", "govern", "adgm"]):
        keywords.extend(["jurisdiction", "governed", "adgm", "law of the abu dhabi global market"])
    if any(k in text for k in ["article", "paragraph", "number"]):
        keywords.extend(["article", "paragraph", "number", "articles of association"])
    if any(k in text for k in ["employment", "contract"]):
        keywords.extend(["employment contract", "employee", "employer"])
    if any(k in text for k in ["register", "members", "directors"]):
        keywords.extend(["register of members", "register of directors", "register"])
    keywords.extend(["abu dhabi global market", "adgm courts"])
    seen = set()
    uniq: List[str] = []
    for k in keywords:
        if k not in seen:
            uniq.append(k)
            seen.add(k)
    return uniq


def _build_comment_text(issue: Dict[str, Any]) -> str:
    parts: List[str] = []
    issue_desc = issue.get("issue") or "Issue"
    parts.append(f"Issue: {issue_desc}")
    severity = issue.get("severity")
    if severity:
        parts.append(f"Severity: {severity}")
    suggestion = issue.get("suggestion")
    if suggestion:
        parts.append(f"Suggestion: {suggestion}")
    return " \n".join(parts)


def _add_inline_marker(paragraph, text: str) -> None:
    short = text
    if len(short) > 140:
        short = short[:140].rstrip() + "…"
    # Add a clear banner-like marker
    run = paragraph.add_run(" ⚠️ [ADGM COMMENT]: ")
    run.bold = True
    run.italic = True
    try:
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    except Exception:
        pass
    # Detail text
    run2 = paragraph.add_run(short)
    run2.italic = True
    try:
        run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    except Exception:
        pass


def add_issue_comments(
    document: Document,
    issues: List[Dict[str, Any]],
    author: str,
    initials: str,
    add_inline_marker: bool = True,
) -> Document:
    paragraph_adds = hasattr(document.paragraphs[0].__class__, "add_comment") if document.paragraphs else False

    for issue in issues:
        if not document.paragraphs:
            break
        anchor_idx: Optional[int] = None
        section = issue.get("section")
        if isinstance(section, str) and section.strip():
            anchor_idx = _find_paragraph_index(document, section)
        if anchor_idx is None:
            for kw in _heuristic_keywords(issue.get("issue") or ""):
                anchor_idx = _find_paragraph_index(document, kw)
                if anchor_idx is not None:
                    break
        if anchor_idx is None:
            anchor_idx = 0

        comment_text = _build_comment_text(issue)
        paragraph = document.paragraphs[anchor_idx]
        try:
            if paragraph_adds:
                paragraph.add_comment(comment_text, author=author, initials=initials)
            else:
                _add_inline_marker(paragraph, comment_text)
        except Exception:
            _add_inline_marker(paragraph, comment_text)

        # Always add a conspicuous inline marker so it won't blend with text
        if add_inline_marker:
            _add_inline_marker(paragraph, issue.get("issue") or "Issue")

    return document


def add_comments_stub(document: Document, issues: List[Dict[str, Any]]) -> Document:
    return document
