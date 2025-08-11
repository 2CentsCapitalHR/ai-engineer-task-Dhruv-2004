from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document


def load_document_from_bytes(data: bytes) -> Document:
    buffer = BytesIO(data)
    return Document(buffer)


def detect_heading_level(style_name: Optional[str]) -> Optional[int]:
    if not style_name:
        return None
    # Common Word heading styles: "Heading 1", "Heading 2", etc.
    if style_name.lower().startswith("heading"):
        parts = style_name.split()
        if len(parts) == 2 and parts[1].isdigit():
            return int(parts[1])
    return None


def parse_document_structure(doc: Document) -> List[Dict[str, Any]]:
    """
    Returns a list of blocks with type and content:
    - heading: {type: "heading", level: int, text: str}
    - paragraph: {type: "paragraph", text: str}
    """
    blocks: List[Dict[str, Any]] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue

        style_name = getattr(paragraph.style, "name", None)
        heading_level = detect_heading_level(style_name)
        if heading_level is not None:
            blocks.append({
                "type": "heading",
                "level": heading_level,
                "text": text,
            })
        else:
            blocks.append({
                "type": "paragraph",
                "text": text,
            })
    return blocks


def summarize_structure(blocks: List[Dict[str, Any]]) -> Dict[str, Any]:
    num_headings = sum(1 for b in blocks if b["type"] == "heading")
    num_paragraphs = sum(1 for b in blocks if b["type"] == "paragraph")
    return {
        "num_headings": num_headings,
        "num_paragraphs": num_paragraphs,
        "preview": blocks[:10],  # preview first 10 blocks
    }


def extract_full_text(doc: Document, max_chars: int = 20000) -> str:
    text = "\n".join(p.text for p in doc.paragraphs)
    return text[:max_chars]
